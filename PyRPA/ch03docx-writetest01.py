import docx

msword = docx.Document() # 새로운 word 파일 생성

msword.add_heading('머릿글',level=0)

par = msword.add_paragraph('첫번째 단락')
par2 = par.insert_paragraph_before('첫번째 단락앞 추가내용')

table = msword.add_table(rows=1,cols=3, style="Table Grid")
table.cell(0,0).text = '첫번째행첫번째열'
table.rows[0].cells[1].text = '첫번째행2번째열'
table.columns[2].cells[0].text = '첫번째행3번째열'

for i in range(10):
    # msword.add_heading('level{} 제목추가'.format(i),level=i)
    msword.add_heading(f'level {i} 제목추가 f문자열 테스트',level=i)

msword.save('./output/docx_write.docx')



