from docx import Document

doc = Document('./output/docx_write.docx')

doc.paragraphs[0].text = doc.paragraphs[0].text+' 수정 추가 내용'

table = doc.tables[0]

new_row = table.add_row()
r_num = len(table.rows)
c_num = len(table.columns)

for i in range(c_num):
    new_row.cells[i].text = f'({r_num}, {i}) 셀추가함'

print(len(doc.paragraphs))

for p in doc.paragraphs:
    print(p.text)

for r in range(r_num):
    for c in range(c_num):
        print(f'({r}, {c})의 내용: '+table.cell(r,c).text)

doc.save('./output/docx_read.docx')        

