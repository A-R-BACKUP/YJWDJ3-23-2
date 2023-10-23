from docx import Document
import docx


doc = Document()

par = doc.add_paragraph() # 빈 단락 생성

for i in range(5):
    par.add_run(f'{i+1} 단락에 {i+1}의 Run객체를 추가')
    r_num = len(par.runs)
    print(f'({i+1}) run 갯수: {r_num}')

r = par.runs
r[0].font.bold = True
r[1].font.italic = True
r[2].font.underline = True
r[3].font.size= docx.shared.Pt(35)
r[4].font.color.rgb = docx.shared.RGBColor(255,0,0)

doc.save('./output/docx_styleTest.docx')
